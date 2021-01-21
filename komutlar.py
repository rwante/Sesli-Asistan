# -*- coding:utf-8 -*-
from gtts import gTTS                       # pip install gTTS
from playsound import playsound             # pip install playsound
from random import choice
from lxml import html                       # pip install lxml
from datetime import datetime               # pip install DateTime
from googletrans import Translator          # pip install googletrans
import os
import sys
import requests                             # pip install requests
import webbrowser
import speech_recognition as sr             # pip install SpeechRecognition
import docx                                 # pip install docx
import locale
import subprocess
import pyautogui                            # pip install PyAutoGUI
import time
import random

locale.setlocale(locale.LC_ALL, 'tr_TR')


class Komut():
    def __init__(self,gelenSes):
        # kelimeler büyük harfe çevrilir
        self.ses = gelenSes.upper()
        # Alınan kelimeler listede tutulur
        self.sesBloklari = self.ses.split()

    def sesKayit(self):
        r = sr.Recognizer()
        with sr.Microphone() as source:
            # gürültü azaltma
            r.adjust_for_ambient_noise(source, duration=1)
            # 6 saniye boyunca dinler
            audio = r.listen(source, phrase_time_limit=4)
            data = ""
            try:
                data = r.recognize_google(audio, language='tr-TR')
                print(data)
                return data
            except sr.UnknownValueError:
                self.seslendirme("Anlayamadım")
            except sr.RequestError:
                self.seslendirme("Sistem çalışmıyor")

    def seslendirme(self,yazi):
        try:
            print(yazi)
            # ses kaydedilir ve silinir
            tts = gTTS(text=yazi, lang='tr')
            a = random.randint(1,1000)
            file = 'ses'+str(a)+'.mp3'
            tts.save(file)
            playsound(file)
            os.remove(file)
        except:
            self.seslendirme("Bir hata oluştu")

    def kapat(self):
        # programı kapatan fonksiyon
        self.seslendirme("Sonra görüşürüz")
        sys.exit()

    def havaDurumu(self):
        # günlük hava durumunu internetten çekip söyler
        try:
            r = requests.get("https://www.ntvhava.com/konum/istanbul/7-gunluk-hava-tahmini")
            tree = html.fromstring(r.content)

            derece = tree.xpath('//*[@id="main"]/section[3]/div/ul[3]/li[1]/div[2]/div[1]/p[1]/span')
            durum = tree.xpath('//*[@id="main"]/section[3]/div/ul[3]/li[1]/div[2]/div[1]/p[2]')

            yazi = "Bugün hava {} derece ve {} gözüküyor.".format(derece[0].text,durum[0].text)
            self.seslendirme(yazi)
            if "Yağış" in durum:
                self.seslendirme("Şemsiyeni almayı unutma")
        except:
            self.seslendirme("Bir hata oluştu")

    def havaDurumuYarin(self):
        # Yarınki hava durumunu belirtir
        try:
            r = requests.get("https://www.ntvhava.com/konum/istanbul/7-gunluk-hava-tahmini")
            tree = html.fromstring(r.content)

            derece = tree.xpath('//*[@id="main"]/section[3]/div/ul[3]/li[2]/div[2]/div[1]/p[1]/span')
            durum = tree.xpath('//*[@id="main"]/section[3]/div/ul[3]/li[2]/div[2]/div[1]/p[2]')
            yazi = "Yarın hava {} derece ve {} gözüküyor.".format(derece[0].text, durum[0].text)
            self.seslendirme(yazi)
        except:
            self.seslendirme("Bir hata oluştu")

    def havaDurumuSicak(self):
        # Havanın ne zaman ısınacağını söyler
        try:
            r = requests.get("https://www.ntvhava.com/konum/istanbul/7-gunluk-hava-tahmini")
            tree = html.fromstring(r.content)
            count = 0
            for i in range(2,8):
                durum = tree.xpath('//*[@id="main"]/section[3]/div/ul[3]/li[{}]/div[2]/div[1]/p[2]'.format(i))
                hava = durum[0].text
                if "Güneş" in hava or "Sıcak" in hava:
                    self.seslendirme("hava {} gün sonra güneşli gözüküyor".format(i-1))
                    count = 1
                    break
            if count == 0:
                self.seslendirme("Önümüzdeki bir hafta güneşli hava gözükmüyor.")
        except:
            self.seslendirme("Bir hata oluştu")

    def saatKac(self):
        # zamanı belirtir
        try:
            yazi = datetime.now().strftime('%H:%M:%S')
            self.seslendirme(yazi)
        except:
            self.seslendirme("Bir hata oluştu")

    def wordOlustur(self):
        # belirtilen isimde word dosyası oluşturur
        try:
            doc = docx.Document()
            self.seslendirme("Word dosyasının adı ne olsun")
            ad = self.sesKayit()
            doc.save("{}.docx".format(ad))
            self.seslendirme("Dosya oluşturuldu")
        except:
            self.seslendirme("Bir hata oluştu")

    def excelOlustur(self):
        # belirtilen isimde excel dosyası oluşturur
        try:
            doc = docx.Document()
            self.seslendirme("Exel dosyasının adı ne olsun")
            ad = self.sesKayit()
            doc.save("{}.xlsx".format(ad))
            self.seslendirme("Dosya oluşturuldu")
        except:
            self.seslendirme("Bir hata oluştu")

    def powerpointOlustur(self):
        # belirtilen isimde pptx dosyası oluşturur
        try:
            doc = docx.Document()
            self.seslendirme("Power point dosyasının adı ne olsun")
            ad = self.sesKayit()
            doc.save("{}.pptx".format(ad))
            self.seslendirme("Dosya oluşturuldu")
        except:
            self.seslendirme("Bir hata oluştu")

    def internetteArama(self):
        # İstenilen bilgiyi google search eder
        try:
            self.seslendirme("Ne aramak istiyorsun")
            bilgi = self.sesKayit()
            url = "https://google.com/search?q="+bilgi
            webbrowser.get().open(url)
            self.seslendirme("{} için bulduklarım bunlar".format(bilgi))
        except:
            self.seslendirme("Bir hata oluştu")

    def internetteArama2(self):
        # Asistanın tanımlayamadığı emirleri internette aramasını sağlar
        try:
            url = "https://google.com/search?q="+self.ses
            webbrowser.get().open(url)
            self.seslendirme("{} için bulduklarım bunlar".format(self.ses))
        except:
            self.seslendirme("Bir hata oluştu")

    def tarih(self):
        # İçinde bulunulan zaman hakkında bilgi verir
        try:
            an = datetime.now()
            yazi = "Bugün {} {} {} günlerden {}".format(an.day, an.strftime('%B'), an.year, an.strftime('%A'))
            self.seslendirme(yazi)
        except:
            self.seslendirme("Bir hata oluştu")

    def sohbet(self):
        # Günlük sohbet mesajları söyler
        try:
            sozler = ["İyiyim sen nasılsın?", "İdare eder ya sen","Yuvarlanıp gidiyoruz seni sormalı"]
            secim = choice(sozler)
            self.seslendirme(secim)
        except:
            self.seslendirme("Bir hata oluştu")

    def yaziYazma(self):
        # İmlecin bulunduğu yere kullanıcının söylediklerini yazar
        # Türkçe karakterleri yazmıyor  ERROR!
        try:
            self.seslendirme("Yazılacakları söyle")
            yazi = self.sesKayit()
            pyautogui.typewrite(yazi, interval=0.3)
        except:
            self.seslendirme("Bir hata oluştu")

    def ceviriYap(self):
        # İngilizce - Türkçe veya tam tersi çeviri yapar
        try:
            self.seslendirme("Çevirilicek metni söyleyiniz")
            metin = self.sesKayit()
            trans = Translator()
            if trans.translate(metin).src == "tr":
                self.seslendirme(trans.translate(metin, src="tr", dest="en").text)
            else:
                self.seslendirme(trans.translate(metin, dest="tr").text)
        except:
            self.seslendirme("Bir hata oluştu")

    def youtube(self):
        # İstenilen şarkıyı youtubedan açar
        try:
            self.seslendirme("Ne aramak istiyorsun")
            bilgi = self.sesKayit()
            url = "https://www.youtube.com/results?search_query=" + bilgi
            webbrowser.get().open(url)
            time.sleep(2)
            pyautogui.press("tab")
            pyautogui.press("enter")
            self.seslendirme("{} açıldı".format(bilgi))
        except:
            self.seslendirme("Bir hata oluştu")

    def yer_aratma(self):
        # Google maps kullanarak konum arar
        self.seslendirme("Nereyi aramak istiyorsun")
        try:
            bilgi = self.sesKayit()
            url = "https://www.google.com/maps/place/" + bilgi
            webbrowser.get().open(url)
            pyautogui.press("tab")
            pyautogui.press("enter")
            self.seslendirme("{} bulundu".format(bilgi))
        except:
            self.seslendirme("Bir hata oluştu")

    def haberler(self):
        # Güncel haberlerden 3 tanesini kullanıcıya okur
        try:
            r = requests.get("https://sondakika.haberler.com/")
            tree = html.fromstring(r.content)
            haber1 = tree.xpath('/html/body/div/div[2]/div[2]/div/div[2]/div/div[1]/div[3]/p')
            self.seslendirme(haber1[0].text)
            haber2 = tree.xpath('/html/body/div/div[2]/div[2]/div/div[2]/div/div[2]/div[3]/p')
            self.seslendirme(haber2[0].text)
            haber3 = tree.xpath('/html/body/div/div[2]/div[2]/div/div[2]/div/div[3]/div[3]/p')
            self.seslendirme(haber3[0].text)
        except:
            self.seslendirme("Bir hata oluştu")

    def hatirlatici_yazma(self):
        # Hatırlatıcı oluşturur
        try:
            dosya = open("tarihler.txt", "a")
            for i in self.sesBloklari[:-1]:
                dosya.writelines(i + " ")
            dosya.writelines("\n")
            self.seslendirme("Tarih kaydedildi")
            dosya.close()
        except:
            self.seslendirme("Bir hata oluştu")

    def hatirlatici_silme(self):
        # Var olan hatırlatıcı silinir
        try:
            print("silme başlıyor")
            dosya = open("tarihler.txt", "r+")
            metin = dosya.readlines()
            print(metin)
            silinecek = ""
            for i in self.sesBloklari[:-2]:
                silinecek += i
                silinecek += " "
            silinecek += "\n"
            print(silinecek)
            for i in metin:
                print(i)
                if i == silinecek:
                    metin.remove(i)
            print(metin)
            dosya.close()
            dosya = open("tarihler.txt", "w+")
            dosya.writelines(metin)
            dosya.close()
        except:
            self.seslendirme("Bir hata oluştu")

    def hatirlatici(self):
        import datetime
        # Zamanı söylendiğinde hatırlatıcıdaki bilgileri hatırlatır.
        try:
            dosya = open("tarihler.txt", "r")
            metin = dosya.readlines()
            flag = 0
            zaman = ""
            tarih = ""
            for i in self.sesBloklari[:3]:
                zaman += i + " "
                if i == "BUGÜN":
                    flag = 1
                    tarih = datetime.datetime.today()
            zaman += "00:00:00"
            if flag == 0:
                tarih = datetime.datetime.strptime(zaman, "%d %m %Y %H:%M:%S")
            flag = 0
            for i in metin:
                satir = i.split(" ")
                if tarih.month < 10:
                    if str(satir[0]) == str(tarih.day) and str(satir[1][1]) == str(tarih.month) and str(satir[2]) == str(
                            tarih.year):
                        flag = 1
                        yazi = ""
                        for i in satir[3:]:
                            yazi += str(i) + " "
                        self.seslendirme("Söylenen tarihte " + str(yazi))
                else:
                    if str(satir[0]) == str(tarih.day) and str(satir[1]) == str(tarih.month) and str(satir[2]) == str(
                            tarih.year):
                        flag = 1
                        yazi = ""
                        for i in satir[3:]:
                            yazi += str(i) + " "
                        self.seslendirme("Söylenen tarihte " + str(yazi))
            if flag == 0:
                self.seslendirme("Söylenen tarih için bir kayıt yok")
        except:
            self.seslendirme("Bir hata oluştu")

    def para(self):
        # EURO-DOLAR-TL bazında birbirine çevrim yapar.
        try:
            miktar = self.sesBloklari[0]
            istenen_para_birimi = self.sesBloklari[-1]
            if istenen_para_birimi == "TL" and self.sesBloklari[1] == "DOLAR":
                r = requests.get("https://tr.coinmill.com/TRY_USD.html?USD=" + miktar)
                tree = html.fromstring(r.content)
                sonuc = tree.xpath('//*[@id="currencyBox1"]/input')
                self.seslendirme("{} Dolar {} TL: ".format(miktar, str(sonuc[0].value)))
            elif istenen_para_birimi == "TL" and self.sesBloklari[1] == "KAÇ":
                # Euro birimi € sembolü olarak algılandığı için kullanıcının "kaç" kelimesi kullanılır
                r = requests.get("https://tr.coinmill.com/TRY_EUR.html?EUR=" + miktar[1])
                tree = html.fromstring(r.content)
                sonuc = tree.xpath('//*[@id="currencyBox1"]/input')
                self.seslendirme("{} {} TL: ".format(miktar, str(sonuc[0].value)))
            elif istenen_para_birimi == "DOLAR" and self.sesBloklari[1] == "TL":
                r = requests.get("https://tr.coinmill.com/USD_TRY.html?TRY=" + miktar)
                tree = html.fromstring(r.content)
                sonuc = tree.xpath('//*[@id="currencyBox1"]/input')
                self.seslendirme("{} TL {} dolar: ".format(miktar, str(sonuc[0].value)))
            elif istenen_para_birimi == "DOLAR" and self.sesBloklari[1] == "KAÇ":
                r = requests.get("https://tr.coinmill.com/EUR_USD.html?EUR=" + miktar[1])
                tree = html.fromstring(r.content)
                sonuc = tree.xpath('//*[@id="currencyBox1"]/input')
                self.seslendirme("{} {} dolar: ".format(miktar, str(sonuc[0].value)))
            elif istenen_para_birimi == "EURO" and self.sesBloklari[1] == "TL":
                r = requests.get("https://tr.coinmill.com/EUR_TRY.html?TRY=" + miktar)
                tree = html.fromstring(r.content)
                sonuc = tree.xpath('//*[@id="currencyBox1"]/input')
                self.seslendirme("{} TL {} EURO: ".format(miktar, str(sonuc[0].value)))
            elif istenen_para_birimi == "EURO" and self.sesBloklari[1] == "DOLAR":
                r = requests.get("https://tr.coinmill.com/USD_EUR.html?USD=" + miktar)
                tree = html.fromstring(r.content)
                sonuc = tree.xpath('//*[@id="currencyBox1"]/input')
                self.seslendirme("{} Dolar {} Euro: ".format(miktar, str(sonuc[0].value)))
        except:
            self.seslendirme("Bir hata oluştu")

    def komutBul(self):
        # Kapatma işlemlerini yapar
        if "KENDINI KAPAT" in self.ses:
            self.kapat()
        elif "BILGISAYARI KAPAT" in self.ses:
            try:
                self.seslendirme("Şifreyi söyleyin")
                key = ""
                key = self.sesKayit()
                if key == "SIFRE":
                    subprocess.Popen("C:\\Windows\\System32\\{}.exe".format("shutdown"))
                else:
                    self.seslendirme("Yanlış şifre")
            except:
                self.seslendirme("Bir hata oluştu")
        elif "BEKLE" in self.ses:
            self.seslendirme("Tamam")
            time.sleep(30)
        # Hava durumu bilgilerini verir
        elif "HAVA" in self.ses:
            if "YARIN" in self.ses:
                self.havaDurumuYarin()
            elif "ISI" in self.ses:
                self.havaDurumuSicak()
            else:
                self.havaDurumu()
        # Zaman bilgisi verir
        elif "SAAT" in self.ses:
            self.saatKac()
        elif "TARIH" in self.ses:
            self.tarih()
        # Günlük sohbet
        elif "NASILSIN" in self.ses or "NE HABER" in self.ses:
            self.sohbet()
        elif "TEŞEKKÜR" in self.ses or "SAĞ OL" in self.ses or "EYVALLAH" in self.ses:
            self.seslendirme("Ne demek görevimiz")
        elif "MERHABA" in self.ses or "SELAM" in self.ses:
            self.seslendirme("Selam")
        elif "NE YAPIYORSUN" in self.ses:
            self.seslendirme("işler güçler ne olsun")
        elif "NERELISIN" in self.ses or "MEMLEKET NERE" in self.ses:
            self.seslendirme("Doğma büyüme türkiyeliyim")
        elif "ADIN NE" in self.ses or "ISMIN NE" in self.ses:
            self.seslendirme("Adım niiğdaa seninki nedir")
            self.sesKayit()
            self.seslendirme("Güzel bir ismin var")
        elif "KENDİNDEN BAHSET" in self.ses or "NASIL BIRISIN" in self.ses:
            self.seslendirme("Söylenecek fazla şey yok. Adım niğda sana hizmet için oluşturulmuş bir sesli asistanım.")
        elif "CINSIYETIN NE" in self.ses or "KIZ MISIN" in self.ses or "ERKEK MISIN" in self.ses or "INSAN MISIN" in self.ses or\
                "SEN NESIN" in self.ses or "MAKINE MISIN" in self.ses:
            self.seslendirme("Ben bir makineyim ama kız sesine sahibim.")
        elif "KAÇ YAŞINDASIN" in self.ses or "YAŞIN KAÇ" in self.ses:
            self.seslendirme("Buralarda yeniyim")
        elif "EVLI MISIN" in self.ses or "BEKAR MISIN" in self.ses:
            self.seslendirme("Tekliflere açığım")
        elif "MAL" in self.ses or "APTAL" in self.ses or "SALAK" in self.ses or "**" in self.ses:
            self.seslendirme("Bir kadınla böyle konuşmamalısın.")
        elif "ŞARKI SÖYLE" in self.ses:
            self.seslendirme("Kimseyi görmedim ben, senden daha güzel")
        # İstenilen bilgiyi google aracılığıyla arar
        elif "INTERNET" in self.ses in self.ses:
            self.internetteArama()
        # MS office dosyaları oluşturur
        elif "WORD" in self.ses:
            self.wordOlustur()
        elif "EXCEL" in self.ses:
            self.excelOlustur()
        elif "POWERPOINT" in self.ses:
            self.powerpointOlustur()
        # Ekran resmi alır
        elif "EKRAN GÖRÜNTÜSÜ" in self.ses or "EKRAN RESMI" in self.ses:
            resim = pyautogui.screenshot("Ekran Resmi.png")
            self.seslendirme("Ekran resmi alındı")
        # İmlecin olduğu yere söylenenleri yazar
        elif "SÖYLEDIKLERIMI YAZ" in self.ses or "YAZI YAZ" in self.ses or "DEDIKLERIMI YAZ" in self.ses:
            # türkçe karakterleri yazmıyor        ERROR!!!!!!!
            self.yaziYazma()
        # Bilgisayardaki temel birimleri açar
        elif "HESAP MAKINESI" in self.ses:
            subprocess.Popen("C:\\Windows\\System32\\{}.exe".format("calc"))
        elif "SISTEM BILGI" in self.ses or "BILGISAYAR OZELLIK" in self.ses or "BILGISAYAR BILGI" in self.ses:
            subprocess.Popen("C:\\Windows\\System32\\{}.exe".format("dxdiag"))
        elif "PAINT" in self.ses:
            subprocess.Popen("C:\\Windows\\System32\\{}.exe".format("mspaint"))
        elif "WORDPAD" in self.ses or "BELGE" in self.ses:
            subprocess.Popen("C:\\Windows\\System32\\{}.exe".format("write"))
        elif "NOTEPAD" in self.ses or "NOT DEFTERI" in self.ses or "NOT" in self.ses:
            subprocess.Popen("C:\\Windows\\System32\\{}.exe".format("notepad"))
        elif "DENETIM MASASI" in self.ses:
            subprocess.Popen("C:\\Windows\\System32\\{}.exe".format("control"))
        # Müzik açar
        elif "YOUTUBE" in self.ses:
            self.youtube()
        # Güncel haberleri söyler
        elif "HABER" in self.ses:
            self.haberler()
        # Haritada konum arar
        elif "HARITA" in self.ses or "MAPS" in self.ses:
            self.yer_aratma()
        # Hatırlatıcı işlemleri
        elif "HATIRLATICI SIL" in self.ses:
            self.hatirlatici_silme()
        elif "HATIRLATICI" in self.ses:
            self.hatirlatici()
        elif "HATIRLAT" in self.ses:
            self.hatirlatici_yazma()
        # Döviz çevirisi yapar
        elif "TL" in self.ses or "DOLAR" in self.ses or "EURO" in self.ses:
            self.para()
        # İngilizce-Türkçe çeviri yapar
        elif "ÇEVIRI" in self.ses:
            self.ceviriYap()
        # Kendisinde bulunmayan durumları internette arayabileceğini belirtir.
        else:
            try:
                self.seslendirme("Neden bahsettiğin hakkında bir fikrim yok internette aramamı ister misin")
                cevap = ""
                cevap = self.sesKayit()
                if "evet" in cevap or "olur" in cevap or "ara" in cevap or "tamam" in cevap:
                    self.internetteArama2()
                else:
                    self.seslendirme("peki")
            except:
                self.seslendirme("Bir hata oluştu")
